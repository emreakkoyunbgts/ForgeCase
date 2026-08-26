"""
PUBLISHER — Ahmet

Case study -> branded BGTS document.

    python -m publisher.publisher <case_study.json> --out out/eng-01.docx
    python -m publisher.publisher <case_study.json> --out out/eng-01.pdf

See the Project Specification, sections 4.1 and 7.
"""
import argparse
import calendar
import hashlib
import json
import sys
from datetime import date
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from reportlab.lib.colors import HexColor
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import (
    KeepInFrame,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

from common.contract import load_seed
from common.errors import die

TEMPLATE = "caseforge-testdata/templates/case_study_template.docx"

NAVY = HexColor("#1B2A4A")
ORANGE = HexColor("#C45C26")

PDF_LAYOUTS = (
    "full-case-study",
    "one-pager",
    "single-slide",
)


def safe_text(value):
    """Return display-ready text, using [MISSING] for blank values."""
    if value is None:
        return "[MISSING]"
    if isinstance(value, str):
        return value if value.strip() else "[MISSING]"
    return str(value)


def anonymise_text(value, real_client, client_type):
    """Replace a real client name in text with the safe client type."""
    if not isinstance(value, str):
        return value
    return value.replace(real_client, client_type)


def print_case_study(case_study):
    """Print a case study as readable JSON."""
    print(json.dumps(case_study, indent=2, ensure_ascii=False))


def collect_provenance_sources(case_study):
    """Collect source records and unique source references from a case study."""
    source_records = []
    engagement_id = case_study.get("engagement_id")
    if isinstance(engagement_id, str) and engagement_id.strip():
        source_records.append(engagement_id.strip())

    source_references = []
    seen = set()
    citations = case_study.get("citations")
    if not isinstance(citations, list):
        citations = []

    for citation in citations:
        if not isinstance(citation, dict):
            continue
        source_ref = citation.get("source_ref")
        if not isinstance(source_ref, str):
            continue
        cleaned = source_ref.strip()
        if not cleaned or cleaned in seen:
            continue
        seen.add(cleaned)
        source_references.append(cleaned)

    return {
        "source_records": source_records,
        "source_references": source_references,
    }


def compute_content_hash(published_content):
    """Return a deterministic SHA-256 hash of safe published content."""
    canonical = json.dumps(
        published_content,
        sort_keys=True,
        separators=(",", ":"),
        ensure_ascii=False,
    )
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


def _subtract_calendar_months(value, months):
    """Return the date that is months calendar months before value."""
    year = value.year
    month = value.month - months
    while month <= 0:
        month += 12
        year -= 1
    day = min(value.day, calendar.monthrange(year, month)[1])
    return date(year, month, day)


def compute_freshness(completed_at, as_of_date):
    """Classify published content freshness from completed_at and as_of_date."""
    if completed_at is None:
        return {
            "freshness_status": "UNKNOWN",
            "freshness_reason": "DATE_MISSING",
        }
    if not isinstance(completed_at, str) or not completed_at.strip():
        return {
            "freshness_status": "UNKNOWN",
            "freshness_reason": "DATE_MISSING",
        }

    try:
        completed_date = date.fromisoformat(completed_at.strip())
        as_of = date.fromisoformat(as_of_date)
    except ValueError:
        return {
            "freshness_status": "UNKNOWN",
            "freshness_reason": "DATE_INVALID",
        }

    cutoff_date = _subtract_calendar_months(as_of, 6)
    if completed_date < cutoff_date:
        return {
            "freshness_status": "STALE",
            "freshness_reason": "OLDER_THAN_SIX_MONTHS",
        }
    return {
        "freshness_status": "FRESH",
        "freshness_reason": "WITHIN_SIX_MONTHS",
    }


def build_provenance_metadata(case_study, published_content, as_of_date):
    """Combine sources, content hash and freshness into one metadata dict."""
    sources = collect_provenance_sources(case_study)
    freshness = compute_freshness(
        case_study.get("completed_at"),
        as_of_date,
    )
    return {
        "source_records": sources["source_records"],
        "source_references": sources["source_references"],
        "completed_at": case_study.get("completed_at"),
        "as_of_date": as_of_date,
        "content_hash": compute_content_hash(published_content),
        "freshness_status": freshness["freshness_status"],
        "freshness_reason": freshness["freshness_reason"],
    }


def write_provenance_sidecar(asset_path, metadata):
    """Write provenance metadata beside a published asset as UTF-8 JSON."""
    asset_path = Path(asset_path)
    sidecar_path = asset_path.with_name(asset_path.name + ".provenance.json")
    with open(sidecar_path, "w", encoding="utf-8") as f:
        json.dump(
            metadata,
            f,
            indent=2,
            ensure_ascii=False,
            sort_keys=True,
        )
        f.write("\n")
    return sidecar_path


def prepare_display_values(case_study):
    """
    Build the safe display values used by both DOCX and PDF output.

    Returns a dict with title, client_type, challenge, approach,
    technology and outcomes. The real client name is never included.
    """
    sections = case_study.get("sections")
    if not isinstance(sections, dict):
        sections = {}

    record = load_seed(case_study.get("engagement_id"))
    client_type = record["client_type"]
    real_client = record["client"]

    sections = {
        key: anonymise_text(
            safe_text(sections.get(key)), real_client, client_type
        )
        for key in [
            "context",
            "challenge",
            "approach",
            "technology",
            "outcomes",
        ]
    }

    return {
        "title": anonymise_text(
            safe_text(case_study.get("title")), real_client, client_type
        ),
        "client_type": client_type,
        "challenge": sections["challenge"],
        "approach": sections["approach"],
        "technology": sections["technology"],
        "outcomes": sections["outcomes"],
    }


def render_docx(case_study, template_path, out_path, as_of_date=None):
    """Fill the template's {{PLACEHOLDERS}} from the case study."""
    display = prepare_display_values(case_study)
    provenance = build_provenance_metadata(
        case_study,
        display,
        _resolve_as_of_date(as_of_date),
    )

    values = {
        "{{TITLE}}":      display["title"],
        "{{CLIENT}}":     display["client_type"],
        "{{DOMAIN}}":     "",
        "{{REGION}}":     "",
        "{{CHALLENGE}}":  display["challenge"],
        "{{APPROACH}}":   display["approach"],
        "{{TECHNOLOGY}}": display["technology"],
        "{{OUTCOMES}}":   display["outcomes"],
    }

    doc = Document(template_path)
    metadata_keys = ("{{CLIENT}}", "{{DOMAIN}}", "{{REGION}}")
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if all(key in run.text for key in metadata_keys):
                metadata_values = [
                    values[key] for key in metadata_keys if values[key]
                ]
                run.text = " · ".join(metadata_values)
                continue

            for key, value in values.items():
                if key in run.text:
                    run.text = run.text.replace(key, value)

    doc.add_paragraph("")
    for line in _provenance_lines(provenance):
        doc.add_paragraph(line)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    write_provenance_sidecar(out_path, provenance)
    return out_path


def _resolve_as_of_date(as_of_date):
    """Return an explicit as_of_date or today's ISO date."""
    return as_of_date or date.today().isoformat()


def _provenance_lines(metadata):
    """Return plain-text provenance lines for PDF and DOCX output."""
    completed_at = metadata.get("completed_at")
    if completed_at is None or (
        isinstance(completed_at, str) and not completed_at.strip()
    ):
        completed_display = "UNKNOWN"
    else:
        completed_display = str(completed_at)

    source_records = ", ".join(metadata.get("source_records") or [])
    source_references = ", ".join(metadata.get("source_references") or [])

    return [
        "Provenance",
        f"Source records: {source_records}",
        f"Source references: {source_references}",
        f"Content hash: {metadata['content_hash']}",
        f"Freshness: {metadata['freshness_status']}",
        f"Reason: {metadata['freshness_reason']}",
        f"Completed at: {completed_display}",
        f"As of date: {metadata['as_of_date']}",
    ]


def _provenance_flowables(metadata, heading_style, body_style):
    """Build visible provenance Paragraph flowables from metadata."""
    lines = _provenance_lines(metadata)
    return [
        Paragraph(escape(lines[0]), heading_style),
        *[Paragraph(escape(line), body_style) for line in lines[1:]],
    ]


def render_pdf(case_study, out_path, layout="full-case-study", as_of_date=None):
    """Create a branded BGTS PDF using the selected layout."""
    if layout not in PDF_LAYOUTS:
        raise ValueError(
            f"unsupported PDF layout '{layout}'; "
            f"use one of: {', '.join(PDF_LAYOUTS)}"
        )

    display = prepare_display_values(case_study)
    provenance = build_provenance_metadata(
        case_study,
        display,
        _resolve_as_of_date(as_of_date),
    )

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    page_size = landscape(A4) if layout == "single-slide" else A4

    styles = getSampleStyleSheet()
    brand = ParagraphStyle(
        "Brand",
        parent=styles["Normal"],
        fontName="Helvetica-Bold",
        fontSize=11,
        textColor=NAVY,
        spaceAfter=12,
    )
    title = ParagraphStyle(
        "CaseTitle",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=18,
        textColor=NAVY,
        spaceAfter=10,
    )
    client = ParagraphStyle(
        "ClientLine",
        parent=styles["Normal"],
        fontName="Helvetica-Oblique",
        fontSize=11,
        textColor=ORANGE,
        spaceAfter=18,
    )
    heading = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=13,
        textColor=NAVY,
        spaceBefore=10,
        spaceAfter=6,
    )
    body = ParagraphStyle(
        "BodyText",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=10,
        leading=14,
        spaceAfter=8,
    )
    footer = ParagraphStyle(
        "Footer",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=9,
        textColor=NAVY,
        spaceBefore=24,
    )
    provenance_heading = ParagraphStyle(
        "ProvenanceHeading",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11,
        textColor=NAVY,
        spaceBefore=10,
        spaceAfter=4,
    )
    provenance_body = ParagraphStyle(
        "ProvenanceBody",
        parent=styles["Normal"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        spaceAfter=2,
    )

    if layout == "one-pager":
        one_brand = ParagraphStyle(
            "OneBrand",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=9,
            textColor=NAVY,
            spaceAfter=4,
        )
        one_title = ParagraphStyle(
            "OneTitle",
            parent=styles["Heading1"],
            fontName="Helvetica-Bold",
            fontSize=12,
            textColor=NAVY,
            spaceAfter=4,
        )
        one_client = ParagraphStyle(
            "OneClient",
            parent=styles["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=8,
            textColor=ORANGE,
            spaceAfter=6,
        )
        one_heading = ParagraphStyle(
            "OneHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=9,
            textColor=NAVY,
            spaceBefore=4,
            spaceAfter=2,
        )
        one_body = ParagraphStyle(
            "OneBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=7,
            leading=9,
            spaceAfter=3,
        )
        one_footer = ParagraphStyle(
            "OneFooter",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=7,
            textColor=NAVY,
            spaceBefore=6,
        )
        one_provenance_heading = ParagraphStyle(
            "OneProvenanceHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=7,
            textColor=NAVY,
            spaceBefore=4,
            spaceAfter=1,
        )
        one_provenance_body = ParagraphStyle(
            "OneProvenanceBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=6,
            leading=7,
            spaceAfter=1,
        )

        document = SimpleDocTemplate(
            str(out_path),
            pagesize=A4,
            leftMargin=28,
            rightMargin=28,
            topMargin=24,
            bottomMargin=24,
        )
        avail_width = A4[0] - 56
        avail_height = A4[1] - 48
        story = [
            KeepInFrame(
                avail_width,
                avail_height,
                [
                    Paragraph(escape("BGTS INTERNATIONAL"), one_brand),
                    Paragraph(escape(display["title"]), one_title),
                    Paragraph(escape(display["client_type"]), one_client),
                    Paragraph(escape("THE CHALLENGE"), one_heading),
                    Paragraph(escape(display["challenge"]), one_body),
                    Paragraph(escape("OUR APPROACH"), one_heading),
                    Paragraph(escape(display["approach"]), one_body),
                    Paragraph(escape("TECHNOLOGY"), one_heading),
                    Paragraph(escape(display["technology"]), one_body),
                    Paragraph(escape("OUTCOMES"), one_heading),
                    Paragraph(escape(display["outcomes"]), one_body),
                    *_provenance_flowables(
                        provenance,
                        one_provenance_heading,
                        one_provenance_body,
                    ),
                    Paragraph(
                        escape("Confidential — BGTS International"),
                        one_footer,
                    ),
                ],
                mode="shrink",
            )
        ]
    elif layout == "single-slide":
        slide_heading = ParagraphStyle(
            "SlideHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11,
            textColor=NAVY,
            spaceBefore=0,
            spaceAfter=4,
        )
        slide_body = ParagraphStyle(
            "SlideBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=8,
            leading=10,
            spaceAfter=0,
        )
        slide_provenance_heading = ParagraphStyle(
            "SlideProvenanceHeading",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=8,
            textColor=NAVY,
            spaceBefore=4,
            spaceAfter=2,
        )
        slide_provenance_body = ParagraphStyle(
            "SlideProvenanceBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=6,
            leading=8,
            spaceAfter=1,
        )

        def slide_cell(heading_text, body_text):
            return [
                Paragraph(escape(heading_text), slide_heading),
                Paragraph(escape(body_text), slide_body),
            ]

        header = [
            Paragraph(escape("BGTS INTERNATIONAL"), brand),
            Paragraph(escape(display["title"]), title),
            Paragraph(escape(display["client_type"]), client),
        ]

        grid = Table(
            [
                [
                    slide_cell("THE CHALLENGE", display["challenge"]),
                    slide_cell("OUR APPROACH", display["approach"]),
                ],
                [
                    slide_cell("TECHNOLOGY", display["technology"]),
                    slide_cell("OUTCOMES", display["outcomes"]),
                ],
            ],
            colWidths=["50%", "50%"],
        )
        grid.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )

        document = SimpleDocTemplate(
            str(out_path),
            pagesize=page_size,
            leftMargin=36,
            rightMargin=36,
            topMargin=28,
            bottomMargin=28,
        )
        avail_width = page_size[0] - 72
        avail_height = page_size[1] - 56
        story = [
            KeepInFrame(
                avail_width,
                avail_height,
                header
                + [
                    grid,
                    Spacer(1, 8),
                    *_provenance_flowables(
                        provenance,
                        slide_provenance_heading,
                        slide_provenance_body,
                    ),
                    Paragraph(
                        escape("Confidential — BGTS International"),
                        footer,
                    ),
                ],
                mode="shrink",
            )
        ]
    else:
        story = [
            Paragraph(escape("BGTS INTERNATIONAL"), brand),
            Paragraph(escape(display["title"]), title),
            Paragraph(escape(display["client_type"]), client),
            Paragraph(escape("The Challenge"), heading),
            Paragraph(escape(display["challenge"]), body),
            Paragraph(escape("Our Approach"), heading),
            Paragraph(escape(display["approach"]), body),
            Paragraph(escape("Technology"), heading),
            Paragraph(escape(display["technology"]), body),
            Paragraph(escape("Outcomes"), heading),
            Paragraph(escape(display["outcomes"]), body),
            Spacer(1, 12),
            *_provenance_flowables(
                provenance,
                provenance_heading,
                provenance_body,
            ),
            Spacer(1, 18),
            Paragraph(escape("Confidential — BGTS International"), footer),
        ]
        document = SimpleDocTemplate(str(out_path), pagesize=page_size)

    document.build(story)
    write_provenance_sidecar(out_path, provenance)
    return out_path


def main():
    parser = argparse.ArgumentParser(
        description="Case study -> branded document"
    )
    parser.add_argument("case_study")
    parser.add_argument(
        "--out",
        default="out/case_study.docx",
    )
    parser.add_argument(
        "--template",
        default=TEMPLATE,
    )
    parser.add_argument(
        "--layout",
        choices=PDF_LAYOUTS,
        default="full-case-study",
    )
    parser.add_argument(
        "--print-json",
        action="store_true",
    )
    parser.add_argument(
        "--as-of-date",
        default=None,
        help=(
            "Freshness evaluation date in YYYY-MM-DD format. "
            "Defaults to the publication date when omitted."
        ),
    )
    args = parser.parse_args()

    try:
        with open(args.case_study, encoding="utf-8") as f:
            case_study = json.load(f)
    except FileNotFoundError:
        die(f"no such file: {args.case_study}")
    except json.JSONDecodeError as e:
        die(f"{args.case_study} is not valid JSON: {e}")

    if args.print_json:
        print_case_study(case_study)
        return

    out_path = Path(args.out)
    suffix = out_path.suffix.lower()

    if suffix == ".docx":
        if args.layout != "full-case-study":
            die(
                "runtime layout selection is available for "
                "PDF output only"
            )

        written = render_docx(
            case_study,
            args.template,
            out_path,
            as_of_date=args.as_of_date,
        )
    elif suffix == ".pdf":
        written = render_pdf(
            case_study,
            out_path,
            layout=args.layout,
            as_of_date=args.as_of_date,
        )
    else:
        die(
            f"unsupported output type '{suffix}' "
            "— use .docx or .pdf"
        )

    print(f"[publisher] wrote {written}", file=sys.stderr)

if __name__ == "__main__":
    main()
