"""Tests for the Publisher."""
from copy import deepcopy
import json
import sys

import pytest
from docx import Document
from common.contract import load_seed
from publisher.publisher import anonymise_text, main, render_docx, render_pdf


def read_docx_text(path):
    """Return all normal paragraph text from a DOCX file."""
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def test_renders_without_crashing(tmp_path):
    with open("caseforge-testdata/case_studies/eng-01_clean.json") as f:
        case_study = json.load(f)

    out = tmp_path / "eng-01.docx"
    written = render_docx(
        case_study,
        "caseforge-testdata/templates/case_study_template.docx",
        out,
    )

    assert written.exists()
    assert written.suffix == ".docx"

    content = read_docx_text(written)
    sections = case_study["sections"]
    record = load_seed(case_study["engagement_id"])
    for expected in [
        case_study["title"],
        sections["challenge"],
        sections["approach"],
        sections["technology"],
        sections["outcomes"],
    ]:
        assert expected in content

    metadata_line = next(
        line for line in content.splitlines() if record["client_type"] in line
    )
    assert not metadata_line.rstrip().endswith("·")
    assert "· ·" not in metadata_line

    for placeholder in [
        "{{TITLE}}",
        "{{CHALLENGE}}",
        "{{APPROACH}}",
        "{{TECHNOLOGY}}",
        "{{OUTCOMES}}",
    ]:
        assert placeholder not in content


def test_real_client_name_never_leaks(tmp_path):
    """eng-01 may NOT be named. Check the real name is nowhere in the output."""
    with open("caseforge-testdata/case_studies/eng-01_clean.json") as f:
        case_study = json.load(f)

    record = load_seed("eng-01")
    out = tmp_path / "eng-01.docx"
    written = render_docx(case_study, "caseforge-testdata/templates/case_study_template.docx", out)

    content = read_docx_text(written)
    assert record["client"] not in content, \
        "the real client name leaked into a published document — spec section 7"


def test_poisoned_client_name_is_anonymised(tmp_path):
    with open("caseforge-testdata/case_studies/eng-01_POISONED.json") as f:
        case_study = json.load(f)

    record = load_seed("eng-01")
    assert record["client"] in case_study["sections"]["context"]
    safe_context = anonymise_text(
        case_study["sections"]["context"],
        record["client"],
        record["client_type"],
    )
    assert record["client"] not in safe_context
    assert record["client_type"] in safe_context

    out = tmp_path / "eng-01-poisoned.docx"
    written = render_docx(
        case_study,
        "caseforge-testdata/templates/case_study_template.docx",
        out,
    )

    content = read_docx_text(written)
    assert record["client"] not in content
    assert record["client_type"] in content


def test_missing_fields_render_as_missing(tmp_path):
    with open("caseforge-testdata/case_studies/eng-01_clean.json") as f:
        original = json.load(f)

    case_study = deepcopy(original)
    case_study["title"] = None
    case_study["sections"]["challenge"] = None
    case_study["sections"]["approach"] = ""
    case_study["sections"]["technology"] = "   "
    del case_study["sections"]["outcomes"]

    out = tmp_path / "eng-01-missing.docx"
    written = render_docx(
        case_study,
        "caseforge-testdata/templates/case_study_template.docx",
        out,
    )

    assert written.exists()
    assert written.suffix == ".docx"

    content = read_docx_text(written)
    assert "[MISSING]" in content
    assert content.count("[MISSING]") >= 5
    assert "None" not in content

    for placeholder in [
        "{{TITLE}}",
        "{{CHALLENGE}}",
        "{{APPROACH}}",
        "{{TECHNOLOGY}}",
        "{{OUTCOMES}}",
    ]:
        assert placeholder not in content


def test_export_to_pdf(tmp_path):
    import pdfplumber

    with open("caseforge-testdata/case_studies/eng-01_clean.json") as f:
        case_study = json.load(f)

    record = load_seed("eng-01")
    out = tmp_path / "eng-01.pdf"
    written = render_pdf(case_study, out)

    assert written.exists()
    assert written.suffix == ".pdf"
    assert written.stat().st_size > 0
    assert written.read_bytes()[:4] == b"%PDF"

    with pdfplumber.open(written) as pdf:
        content = "\n".join((page.extract_text() or "") for page in pdf.pages)

    assert case_study["title"] in content
    assert record["client_type"] in content
    assert record["client"] not in content
    for heading in ["The Challenge", "Our Approach", "Technology", "Outcomes"]:
        assert heading in content


def test_print_json_skips_document_output(tmp_path, monkeypatch, capsys):
    with open("caseforge-testdata/case_studies/eng-01_clean.json", encoding="utf-8") as f:
        case_study = json.load(f)

    docx_out = tmp_path / "should-not-exist.docx"
    pdf_out = tmp_path / "should-not-exist.pdf"
    monkeypatch.setattr(
        sys,
        "argv",
        [
            "publisher",
            "caseforge-testdata/case_studies/eng-01_clean.json",
            "--print-json",
            "--out",
            str(docx_out),
        ],
    )

    main()

    captured = capsys.readouterr()
    printed = json.loads(captured.out)
    assert printed["engagement_id"] == case_study["engagement_id"]
    assert printed["title"] == case_study["title"]
    assert not docx_out.exists()
    assert not pdf_out.exists()


@pytest.mark.parametrize(
    "layout",
    [
        "full-case-study",
        "one-pager",
        "single-slide",
    ],
)
def test_all_pdf_layouts_render_from_same_case_study(tmp_path, layout):
    with open(
        "caseforge-testdata/case_studies/eng-01_clean.json",
        encoding="utf-8",
    ) as f:
        case_study = json.load(f)

    out = tmp_path / f"{layout}.pdf"
    written = render_pdf(case_study, out, layout=layout)

    assert written.exists()
    assert written.suffix == ".pdf"
    assert written.stat().st_size > 0
    assert written.read_bytes()[:4] == b"%PDF"

def test_single_slide_pdf_is_one_page_and_landscape(tmp_path):
    import pdfplumber

    with open(
        "caseforge-testdata/case_studies/eng-01_clean.json",
        encoding="utf-8",
    ) as f:
        case_study = json.load(f)

    out = tmp_path / "single-slide.pdf"
    written = render_pdf(case_study, out, layout="single-slide")

    with pdfplumber.open(written) as pdf:
        assert len(pdf.pages) == 1

        page = pdf.pages[0]
        assert page.width > page.height
    
@pytest.mark.parametrize(
    ("layout", "expect_landscape"),
    [
        ("full-case-study", False),
        ("one-pager", False),
        ("single-slide", True),
    ],
)
def test_cli_selects_each_pdf_layout(
    tmp_path,
    monkeypatch,
    layout,
    expect_landscape,
):
    import pdfplumber

    out = tmp_path / f"{layout}-cli.pdf"

    monkeypatch.setattr(
        sys,
        "argv",
        [
            "publisher",
            "caseforge-testdata/case_studies/eng-01_clean.json",
            "--layout",
            layout,
            "--out",
            str(out),
        ],
    )

    main()

    assert out.exists()
    assert out.suffix == ".pdf"
    assert out.stat().st_size > 0
    assert out.read_bytes()[:4] == b"%PDF"

    with pdfplumber.open(out) as pdf:
        assert len(pdf.pages) >= 1
        page = pdf.pages[0]

        if expect_landscape:
            assert page.width > page.height
        else:
            assert page.height > page.width


def test_single_slide_pdf_uses_slide_headings(tmp_path):
    import pdfplumber

    with open(
        "caseforge-testdata/case_studies/eng-01_clean.json",
        encoding="utf-8",
    ) as f:
        case_study = json.load(f)

    out = tmp_path / "single-slide-headings.pdf"
    written = render_pdf(
        case_study,
        out,
        layout="single-slide",
    )

    with pdfplumber.open(written) as pdf:
        content = "\n".join(
            (page.extract_text() or "")
            for page in pdf.pages
        )

    for heading in [
        "THE CHALLENGE",
        "OUR APPROACH",
        "TECHNOLOGY",
        "OUTCOMES",
    ]:
        assert heading in content


def test_one_pager_compacts_long_case_study_to_one_portrait_page(tmp_path):
    import pdfplumber

    with open(
        "caseforge-testdata/case_studies/eng-01_clean.json",
        encoding="utf-8",
    ) as f:
        original = json.load(f)

    case_study = deepcopy(original)

    for section_name in [
        "challenge",
        "approach",
        "technology",
        "outcomes",
    ]:
        original_text = case_study["sections"][section_name]
        case_study["sections"][section_name] = " ".join(
            [original_text] * 12
        )

    full_out = tmp_path / "long-full-case-study.pdf"
    one_pager_out = tmp_path / "long-one-pager.pdf"

    render_pdf(
        case_study,
        full_out,
        layout="full-case-study",
    )
    render_pdf(
        case_study,
        one_pager_out,
        layout="one-pager",
    )

    with pdfplumber.open(full_out) as full_pdf:
        full_page_count = len(full_pdf.pages)

    with pdfplumber.open(one_pager_out) as one_pager_pdf:
        one_pager_page_count = len(one_pager_pdf.pages)
        one_pager_page = one_pager_pdf.pages[0]

    assert full_page_count > 1
    assert one_pager_page_count == 1
    assert one_pager_page.height > one_pager_page.width


def test_cli_rejects_non_default_layout_for_docx(
    tmp_path,
    monkeypatch,
    capsys,
):
    out = tmp_path / "unsupported-layout.docx"

    monkeypatch.setattr(
        sys,
        "argv",
        [
            "publisher",
            "caseforge-testdata/case_studies/eng-01_clean.json",
            "--layout",
            "single-slide",
            "--out",
            str(out),
        ],
    )

    with pytest.raises(SystemExit):
        main()

    captured = capsys.readouterr()

    assert not out.exists()
    assert "PDF output only" in captured.err


def test_collects_provenance_sources_from_case_study():
    import publisher.publisher as publisher_module

    case_study = {
        "engagement_id": "eng-01",
        "citations": [
            {
                "claim": "Payment latency reduced",
                "source_ref": "closeout.pdf#page=5",
            }
        ],
    }

    result = publisher_module.collect_provenance_sources(case_study)

    assert result == {
        "source_records": ["eng-01"],
        "source_references": ["closeout.pdf#page=5"],
    }


def test_provenance_sources_ignore_blank_and_duplicate_references():
    import publisher.publisher as publisher_module

    case_study = {
        "engagement_id": "eng-01",
        "citations": [
            {"claim": "A", "source_ref": "report.pdf#page=1"},
            {"claim": "B", "source_ref": ""},
            {"claim": "C", "source_ref": None},
            {"claim": "D", "source_ref": "report.pdf#page=1"},
            {"claim": "E", "source_ref": "report.pdf#page=2"},
        ],
    }

    result = publisher_module.collect_provenance_sources(case_study)

    assert result["source_records"] == ["eng-01"]
    assert result["source_references"] == [
        "report.pdf#page=1",
        "report.pdf#page=2",
    ]


def test_content_hash_is_independent_of_dictionary_key_order():
    import publisher.publisher as publisher_module

    published_content_a = {
        "title": "Payment Platform Modernisation",
        "client_type": "Fintech company",
        "challenge": "High payment latency",
        "approach": "Rebuilt the processing pipeline",
        "technology": "Python, PostgreSQL",
        "outcomes": "Latency reduced",
    }
    published_content_b = {
        "outcomes": "Latency reduced",
        "technology": "Python, PostgreSQL",
        "approach": "Rebuilt the processing pipeline",
        "challenge": "High payment latency",
        "client_type": "Fintech company",
        "title": "Payment Platform Modernisation",
    }

    hash_a = publisher_module.compute_content_hash(published_content_a)
    hash_b = publisher_module.compute_content_hash(published_content_b)

    assert hash_a == hash_b


def test_content_hash_changes_when_published_content_changes():
    import publisher.publisher as publisher_module

    published_content_a = {
        "title": "Payment Platform Modernisation",
        "client_type": "Fintech company",
        "challenge": "High payment latency",
        "approach": "Rebuilt the processing pipeline",
        "technology": "Python, PostgreSQL",
        "outcomes": "Latency reduced",
    }
    published_content_b = {
        "title": "Payment Platform Modernisation",
        "client_type": "Fintech company",
        "challenge": "High payment latency",
        "approach": "Rebuilt the processing pipeline",
        "technology": "Python, PostgreSQL",
        "outcomes": "Latency reduced by 45 percent",
    }

    hash_a = publisher_module.compute_content_hash(published_content_a)
    hash_b = publisher_module.compute_content_hash(published_content_b)

    assert hash_a != hash_b


def test_content_hash_has_sha256_hex_format():
    import publisher.publisher as publisher_module
    import re

    published_content = {
        "title": "Payment Platform Modernisation",
        "client_type": "Fintech company",
        "challenge": "High payment latency",
        "approach": "Rebuilt the processing pipeline",
        "technology": "Python, PostgreSQL",
        "outcomes": "Latency reduced",
    }

    result = publisher_module.compute_content_hash(published_content)

    assert isinstance(result, str)
    assert len(result) == 64
    assert re.fullmatch(r"[0-9a-f]{64}", result) is not None
    int(result, 16)


@pytest.mark.parametrize(
    "completed_at",
    [
        None,
        "",
        "   ",
    ],
)
def test_freshness_is_unknown_when_completed_at_is_missing(completed_at):
    import publisher.publisher as publisher_module

    result = publisher_module.compute_freshness(
        completed_at,
        "2026-08-06",
    )

    assert result == {
        "freshness_status": "UNKNOWN",
        "freshness_reason": "DATE_MISSING",
    }


def test_freshness_is_unknown_when_completed_at_is_invalid():
    import publisher.publisher as publisher_module

    result = publisher_module.compute_freshness(
        "2026-13-40",
        "2026-08-06",
    )

    assert result == {
        "freshness_status": "UNKNOWN",
        "freshness_reason": "DATE_INVALID",
    }


def test_recent_content_is_fresh():
    import publisher.publisher as publisher_module

    result = publisher_module.compute_freshness(
        "2026-06-01",
        "2026-08-06",
    )

    assert result == {
        "freshness_status": "FRESH",
        "freshness_reason": "WITHIN_SIX_MONTHS",
    }


def test_content_older_than_six_calendar_months_is_stale():
    import publisher.publisher as publisher_module

    result = publisher_module.compute_freshness(
        "2026-02-05",
        "2026-08-06",
    )

    assert result == {
        "freshness_status": "STALE",
        "freshness_reason": "OLDER_THAN_SIX_MONTHS",
    }


def test_content_exactly_six_calendar_months_old_is_fresh():
    import publisher.publisher as publisher_module

    result = publisher_module.compute_freshness(
        "2026-02-06",
        "2026-08-06",
    )

    assert result == {
        "freshness_status": "FRESH",
        "freshness_reason": "WITHIN_SIX_MONTHS",
    }


def test_freshness_uses_calendar_months_not_fixed_180_days():
    import publisher.publisher as publisher_module

    result = publisher_module.compute_freshness(
        "2026-02-28",
        "2026-08-31",
    )

    assert result == {
        "freshness_status": "FRESH",
        "freshness_reason": "WITHIN_SIX_MONTHS",
    }


def test_freshness_handles_leap_day():
    import publisher.publisher as publisher_module

    result = publisher_module.compute_freshness(
        "2024-02-29",
        "2024-08-31",
    )

    assert result == {
        "freshness_status": "FRESH",
        "freshness_reason": "WITHIN_SIX_MONTHS",
    }


def test_build_provenance_metadata_combines_sources_hash_and_missing_date():
    import publisher.publisher as publisher_module

    case_study = {
        "engagement_id": "eng-01",
        "citations": [
            {
                "claim": "Payment latency reduced",
                "source_ref": "closeout.pdf#page=5",
            }
        ],
    }
    published_content = {
        "title": "Payment Platform Modernisation",
        "client_type": "Fintech company",
        "challenge": "High payment latency",
        "approach": "Rebuilt the processing pipeline",
        "technology": "Python, PostgreSQL",
        "outcomes": "Latency reduced",
    }
    as_of_date = "2026-08-06"

    result = publisher_module.build_provenance_metadata(
        case_study,
        published_content,
        as_of_date,
    )

    assert result["source_records"] == ["eng-01"]
    assert result["source_references"] == ["closeout.pdf#page=5"]
    assert result["completed_at"] is None
    assert result["as_of_date"] == "2026-08-06"
    assert result["content_hash"] == publisher_module.compute_content_hash(
        published_content
    )
    assert result["freshness_status"] == "UNKNOWN"
    assert result["freshness_reason"] == "DATE_MISSING"


def test_build_provenance_metadata_reports_stale_content():
    import publisher.publisher as publisher_module
    import re

    case_study = {
        "engagement_id": "test-stale",
        "completed_at": "2026-02-05",
        "citations": [
            {
                "claim": "Test claim",
                "source_ref": "test-report.pdf#page=1",
            }
        ],
    }
    published_content = {
        "title": "Stale Content Fixture",
        "client_type": "Regional bank",
        "challenge": "Outdated reporting process",
        "approach": "Automated monthly reporting",
        "technology": "Python",
        "outcomes": "Reporting time reduced",
    }
    as_of_date = "2026-08-06"

    result = publisher_module.build_provenance_metadata(
        case_study,
        published_content,
        as_of_date,
    )

    assert result["completed_at"] == "2026-02-05"
    assert result["as_of_date"] == "2026-08-06"
    assert result["freshness_status"] == "STALE"
    assert result["freshness_reason"] == "OLDER_THAN_SIX_MONTHS"
    assert result["source_records"] == ["test-stale"]
    assert isinstance(result["content_hash"], str)
    assert len(result["content_hash"]) == 64
    assert re.fullmatch(r"[0-9a-f]{64}", result["content_hash"]) is not None


def test_build_provenance_metadata_does_not_mutate_inputs():
    import publisher.publisher as publisher_module

    case_study = {
        "engagement_id": "test-immutable",
        "completed_at": "2026-06-01",
        "citations": [
            {
                "claim": "Immutable fixture claim",
                "source_ref": "immutable-report.pdf#page=3",
            }
        ],
    }
    published_content = {
        "title": "Immutable Content Fixture",
        "client_type": "Fintech company",
        "challenge": "Process drift",
        "approach": "Locked metadata assembly",
        "technology": "Python",
        "outcomes": "Inputs remain unchanged",
    }
    as_of_date = "2026-08-06"

    case_study_before = deepcopy(case_study)
    published_content_before = deepcopy(published_content)

    publisher_module.build_provenance_metadata(
        case_study,
        published_content,
        as_of_date,
    )

    assert case_study == case_study_before
    assert published_content == published_content_before


@pytest.mark.parametrize(
    "layout",
    [
        "full-case-study",
        "one-pager",
        "single-slide",
    ],
)
def test_pdf_layouts_include_visible_provenance_block(tmp_path, layout):
    import pdfplumber
    import re
    import publisher.publisher as publisher_module

    with open(
        "caseforge-testdata/case_studies/eng-01_clean.json",
        encoding="utf-8",
    ) as f:
        case_study = json.load(f)

    out = tmp_path / f"{layout}-provenance.pdf"
    written = publisher_module.render_pdf(
        case_study,
        out,
        layout=layout,
        as_of_date="2026-08-06",
    )

    with pdfplumber.open(written) as pdf:
        page_count = len(pdf.pages)
        content = "\n".join(
            (page.extract_text() or "")
            for page in pdf.pages
        )

    normalized = re.sub(r"\s+", " ", content)

    assert "Provenance" in normalized
    assert "Source records" in normalized
    assert "eng-01" in normalized
    assert "Content hash" in normalized
    assert re.search(
        r"Content hash[:\s]+([0-9a-f]{64})",
        normalized,
    ) is not None
    assert "Freshness" in normalized
    assert "UNKNOWN" in normalized
    assert "DATE_MISSING" in normalized
    assert "As of date" in normalized
    assert "2026-08-06" in normalized
    assert "closeout.pdf#page=5" in normalized

    if layout in ("one-pager", "single-slide"):
        assert page_count == 1


def test_docx_includes_visible_provenance_block(tmp_path):
    import re
    import publisher.publisher as publisher_module

    with open(
        "caseforge-testdata/case_studies/eng-01_clean.json",
        encoding="utf-8",
    ) as f:
        case_study = json.load(f)

    record = load_seed(case_study["engagement_id"])
    out = tmp_path / "eng-01-provenance.docx"
    written = publisher_module.render_docx(
        case_study,
        "caseforge-testdata/templates/case_study_template.docx",
        out,
        as_of_date="2026-08-06",
    )

    document = Document(written)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    normalized = re.sub(r"\s+", " ", "\n".join(parts))

    assert "Provenance" in normalized
    assert "Source records" in normalized
    assert "eng-01" in normalized
    assert "Source references" in normalized
    assert "closeout.pdf#page=5" in normalized
    assert "Content hash" in normalized
    assert re.search(
        r"Content hash[:\s]+([0-9a-f]{64})",
        normalized,
    ) is not None
    assert "Freshness" in normalized
    assert "UNKNOWN" in normalized
    assert "DATE_MISSING" in normalized
    assert "Completed at" in normalized
    assert "As of date" in normalized
    assert "2026-08-06" in normalized
    assert record["client"] not in normalized


def test_write_provenance_sidecar_creates_machine_readable_json(tmp_path):
    import re
    import publisher.publisher as publisher_module

    case_study = {
        "engagement_id": "eng-01",
        "citations": [
            {
                "claim": "Payment latency reduced",
                "source_ref": "closeout.pdf#page=5",
            }
        ],
    }
    published_content = {
        "title": "Payment Platform Modernisation",
        "client_type": "Fintech company",
        "challenge": "High payment latency",
        "approach": "Rebuilt the processing pipeline",
        "technology": "Python, PostgreSQL",
        "outcomes": "Latency reduced",
    }
    metadata = publisher_module.build_provenance_metadata(
        case_study,
        published_content,
        "2026-08-06",
    )
    asset_path = tmp_path / "eng-01.pdf"

    sidecar_path = publisher_module.write_provenance_sidecar(
        asset_path,
        metadata,
    )

    assert sidecar_path == tmp_path / "eng-01.pdf.provenance.json"
    assert sidecar_path.exists()

    with open(sidecar_path, encoding="utf-8") as f:
        loaded = json.load(f)

    assert loaded == metadata
    assert loaded["freshness_status"] == "UNKNOWN"
    assert loaded["freshness_reason"] == "DATE_MISSING"
    assert loaded["completed_at"] is None
    assert loaded["source_records"] == ["eng-01"]
    assert loaded["source_references"] == ["closeout.pdf#page=5"]
    assert isinstance(loaded["content_hash"], str)
    assert len(loaded["content_hash"]) == 64
    assert re.fullmatch(r"[0-9a-f]{64}", loaded["content_hash"]) is not None


def test_sidecar_name_keeps_asset_extension_to_avoid_collisions(tmp_path):
    import publisher.publisher as publisher_module

    case_study = {
        "engagement_id": "eng-01",
        "citations": [
            {
                "claim": "Payment latency reduced",
                "source_ref": "closeout.pdf#page=5",
            }
        ],
    }
    published_content = {
        "title": "Payment Platform Modernisation",
        "client_type": "Fintech company",
        "challenge": "High payment latency",
        "approach": "Rebuilt the processing pipeline",
        "technology": "Python, PostgreSQL",
        "outcomes": "Latency reduced",
    }
    metadata = publisher_module.build_provenance_metadata(
        case_study,
        published_content,
        "2026-08-06",
    )

    pdf_sidecar = publisher_module.write_provenance_sidecar(
        tmp_path / "eng-01.pdf",
        metadata,
    )
    docx_sidecar = publisher_module.write_provenance_sidecar(
        tmp_path / "eng-01.docx",
        metadata,
    )

    assert pdf_sidecar == tmp_path / "eng-01.pdf.provenance.json"
    assert docx_sidecar == tmp_path / "eng-01.docx.provenance.json"
    assert pdf_sidecar != docx_sidecar
    assert pdf_sidecar.exists()
    assert docx_sidecar.exists()


def test_write_provenance_sidecar_does_not_mutate_metadata(tmp_path):
    import publisher.publisher as publisher_module

    case_study = {
        "engagement_id": "test-immutable-sidecar",
        "completed_at": "2026-06-01",
        "citations": [
            {
                "claim": "Immutable sidecar claim",
                "source_ref": "immutable-sidecar.pdf#page=1",
            }
        ],
    }
    published_content = {
        "title": "Immutable Sidecar Fixture",
        "client_type": "Fintech company",
        "challenge": "Metadata mutation risk",
        "approach": "Write without side effects",
        "technology": "Python",
        "outcomes": "Metadata remains unchanged",
    }
    metadata = publisher_module.build_provenance_metadata(
        case_study,
        published_content,
        "2026-08-06",
    )
    metadata_before = deepcopy(metadata)

    publisher_module.write_provenance_sidecar(
        tmp_path / "eng-01.pdf",
        metadata,
    )

    assert metadata == metadata_before


@pytest.mark.parametrize(
    "layout",
    [
        "full-case-study",
        "one-pager",
        "single-slide",
    ],
)
def test_render_pdf_automatically_writes_provenance_sidecar(tmp_path, layout):
    import re
    import publisher.publisher as publisher_module

    with open(
        "caseforge-testdata/case_studies/eng-01_clean.json",
        encoding="utf-8",
    ) as f:
        case_study = json.load(f)

    record = load_seed(case_study["engagement_id"])
    pdf_path = tmp_path / f"{layout}.pdf"
    written = publisher_module.render_pdf(
        case_study,
        pdf_path,
        layout=layout,
        as_of_date="2026-08-06",
    )

    assert written.exists()
    sidecar_path = pdf_path.with_name(pdf_path.name + ".provenance.json")
    assert sidecar_path.exists()

    with open(sidecar_path, encoding="utf-8") as f:
        loaded = json.load(f)

    assert loaded["source_records"] == ["eng-01"]
    assert "closeout.pdf#page=5" in loaded["source_references"]
    assert loaded["completed_at"] is None
    assert loaded["as_of_date"] == "2026-08-06"
    assert loaded["freshness_status"] == "UNKNOWN"
    assert loaded["freshness_reason"] == "DATE_MISSING"
    assert isinstance(loaded["content_hash"], str)
    assert len(loaded["content_hash"]) == 64
    assert re.fullmatch(r"[0-9a-f]{64}", loaded["content_hash"]) is not None
    assert record["client"] not in json.dumps(loaded)

    if layout == "full-case-study":
        one_pager_path = tmp_path / "one-pager-compare.pdf"
        single_slide_path = tmp_path / "single-slide-compare.pdf"
        publisher_module.render_pdf(
            case_study,
            one_pager_path,
            layout="one-pager",
            as_of_date="2026-08-06",
        )
        publisher_module.render_pdf(
            case_study,
            single_slide_path,
            layout="single-slide",
            as_of_date="2026-08-06",
        )
        with open(
            one_pager_path.with_name(
                one_pager_path.name + ".provenance.json"
            ),
            encoding="utf-8",
        ) as f:
            one_pager_meta = json.load(f)
        with open(
            single_slide_path.with_name(
                single_slide_path.name + ".provenance.json"
            ),
            encoding="utf-8",
        ) as f:
            single_slide_meta = json.load(f)
        assert loaded["content_hash"] == one_pager_meta["content_hash"]
        assert loaded["content_hash"] == single_slide_meta["content_hash"]


def test_render_docx_automatically_writes_provenance_sidecar(tmp_path):
    import re
    import publisher.publisher as publisher_module

    with open(
        "caseforge-testdata/case_studies/eng-01_clean.json",
        encoding="utf-8",
    ) as f:
        case_study = json.load(f)

    record = load_seed(case_study["engagement_id"])
    docx_path = tmp_path / "eng-01.docx"
    written = publisher_module.render_docx(
        case_study,
        "caseforge-testdata/templates/case_study_template.docx",
        docx_path,
        as_of_date="2026-08-06",
    )

    assert written.exists()
    sidecar_path = tmp_path / "eng-01.docx.provenance.json"
    assert sidecar_path.exists()

    with open(sidecar_path, encoding="utf-8") as f:
        loaded = json.load(f)

    assert loaded["source_records"] == ["eng-01"]
    assert "closeout.pdf#page=5" in loaded["source_references"]
    assert loaded["completed_at"] is None
    assert loaded["as_of_date"] == "2026-08-06"
    assert loaded["freshness_status"] == "UNKNOWN"
    assert loaded["freshness_reason"] == "DATE_MISSING"
    assert isinstance(loaded["content_hash"], str)
    assert len(loaded["content_hash"]) == 64
    assert re.fullmatch(r"[0-9a-f]{64}", loaded["content_hash"]) is not None
    assert record["client"] not in json.dumps(loaded)


def test_pdf_and_docx_sidecars_use_identical_provenance_metadata(tmp_path):
    import publisher.publisher as publisher_module

    with open(
        "caseforge-testdata/case_studies/eng-01_clean.json",
        encoding="utf-8",
    ) as f:
        case_study = json.load(f)

    pdf_path = tmp_path / "eng-01.pdf"
    docx_path = tmp_path / "eng-01.docx"

    publisher_module.render_pdf(
        case_study,
        pdf_path,
        layout="full-case-study",
        as_of_date="2026-08-06",
    )
    publisher_module.render_docx(
        case_study,
        "caseforge-testdata/templates/case_study_template.docx",
        docx_path,
        as_of_date="2026-08-06",
    )

    pdf_sidecar = pdf_path.with_name(pdf_path.name + ".provenance.json")
    docx_sidecar = docx_path.with_name(docx_path.name + ".provenance.json")

    with open(pdf_sidecar, encoding="utf-8") as f:
        pdf_metadata = json.load(f)
    with open(docx_sidecar, encoding="utf-8") as f:
        docx_metadata = json.load(f)

    assert pdf_metadata == docx_metadata


def test_cli_forwards_as_of_date_to_pdf_renderer(tmp_path, monkeypatch):
    import publisher.publisher as publisher_module

    out = tmp_path / "eng-01-as-of.pdf"
    captured = {}

    def fake_render_pdf(
        case_study,
        out_path,
        layout="full-case-study",
        as_of_date=None,
    ):
        captured["as_of_date"] = as_of_date
        captured["layout"] = layout
        return out_path

    monkeypatch.setattr(publisher_module, "render_pdf", fake_render_pdf)
    monkeypatch.setattr(
        sys,
        "argv",
        [
            "publisher",
            "caseforge-testdata/case_studies/eng-01_clean.json",
            "--layout",
            "full-case-study",
            "--out",
            str(out),
            "--as-of-date",
            "2026-08-06",
        ],
    )

    main()

    assert captured["as_of_date"] == "2026-08-06"
    assert captured["layout"] == "full-case-study"


def test_cli_forwards_as_of_date_to_docx_renderer(tmp_path, monkeypatch):
    import publisher.publisher as publisher_module

    out = tmp_path / "eng-01-as-of.docx"
    captured = {}

    def fake_render_docx(case_study, template_path, out_path, as_of_date=None):
        captured["as_of_date"] = as_of_date
        captured["template_path"] = template_path
        return out_path

    monkeypatch.setattr(publisher_module, "render_docx", fake_render_docx)
    monkeypatch.setattr(
        sys,
        "argv",
        [
            "publisher",
            "caseforge-testdata/case_studies/eng-01_clean.json",
            "--out",
            str(out),
            "--as-of-date",
            "2026-08-06",
        ],
    )

    main()

    assert captured["as_of_date"] == "2026-08-06"


def test_cli_without_as_of_date_remains_supported(tmp_path, monkeypatch):
    import publisher.publisher as publisher_module

    out = tmp_path / "eng-01-default.pdf"
    captured = {}

    def fake_render_pdf(
        case_study,
        out_path,
        layout="full-case-study",
        as_of_date=None,
    ):
        captured["as_of_date"] = as_of_date
        captured["layout"] = layout
        return out_path

    monkeypatch.setattr(publisher_module, "render_pdf", fake_render_pdf)
    monkeypatch.setattr(
        sys,
        "argv",
        [
            "publisher",
            "caseforge-testdata/case_studies/eng-01_clean.json",
            "--layout",
            "one-pager",
            "--out",
            str(out),
        ],
    )

    main()

    assert captured["as_of_date"] is None
    assert captured["layout"] == "one-pager"


@pytest.mark.parametrize(
    "layout",
    [
        "full-case-study",
        "one-pager",
        "single-slide",
    ],
)
def test_pdf_layouts_flag_stale_completed_at_in_asset_and_sidecar(
    tmp_path,
    layout,
):
    """Synthetic completed_at verifies STALE surfaces in PDF asset + sidecar."""
    import pdfplumber
    import re
    import publisher.publisher as publisher_module

    with open(
        "caseforge-testdata/case_studies/eng-01_clean.json",
        encoding="utf-8",
    ) as f:
        case_study = deepcopy(json.load(f))

    # Synthetic acceptance dates only — not a real engagement timeline.
    case_study["completed_at"] = "2026-02-05"
    as_of_date = "2026-08-06"

    record = load_seed(case_study["engagement_id"])
    out = tmp_path / f"{layout}-stale.pdf"
    written = publisher_module.render_pdf(
        case_study,
        out,
        layout=layout,
        as_of_date=as_of_date,
    )

    assert written.exists()
    assert written.suffix == ".pdf"
    assert written.read_bytes()[:4] == b"%PDF"

    with pdfplumber.open(written) as pdf:
        page_count = len(pdf.pages)
        content = "\n".join(
            (page.extract_text() or "")
            for page in pdf.pages
        )

    normalized = re.sub(r"\s+", " ", content)

    assert "Freshness" in normalized
    assert "STALE" in normalized
    assert "OLDER_THAN_SIX_MONTHS" in normalized
    assert "Completed at" in normalized
    assert "2026-02-05" in normalized
    assert "As of date" in normalized
    assert "2026-08-06" in normalized
    assert record["client"] not in normalized

    sidecar_path = out.with_name(out.name + ".provenance.json")
    assert sidecar_path.exists()
    with open(sidecar_path, encoding="utf-8") as f:
        loaded = json.load(f)

    assert loaded["completed_at"] == "2026-02-05"
    assert loaded["as_of_date"] == "2026-08-06"
    assert loaded["freshness_status"] == "STALE"
    assert loaded["freshness_reason"] == "OLDER_THAN_SIX_MONTHS"
    assert isinstance(loaded["content_hash"], str)
    assert len(loaded["content_hash"]) == 64
    assert re.fullmatch(r"[0-9a-f]{64}", loaded["content_hash"]) is not None
    assert record["client"] not in json.dumps(loaded)

    if layout in ("one-pager", "single-slide"):
        assert page_count == 1


def test_docx_flags_stale_completed_at_in_asset_and_sidecar(tmp_path):
    """Synthetic completed_at verifies STALE surfaces in DOCX asset + sidecar."""
    import re
    import publisher.publisher as publisher_module

    with open(
        "caseforge-testdata/case_studies/eng-01_clean.json",
        encoding="utf-8",
    ) as f:
        case_study = deepcopy(json.load(f))

    # Synthetic acceptance dates only — not a real engagement timeline.
    case_study["completed_at"] = "2026-02-05"
    as_of_date = "2026-08-06"

    record = load_seed(case_study["engagement_id"])
    out = tmp_path / "eng-01-stale.docx"
    written = publisher_module.render_docx(
        case_study,
        "caseforge-testdata/templates/case_study_template.docx",
        out,
        as_of_date=as_of_date,
    )

    assert written.exists()

    document = Document(written)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    normalized = re.sub(r"\s+", " ", "\n".join(parts))

    assert "Freshness" in normalized
    assert "STALE" in normalized
    assert "OLDER_THAN_SIX_MONTHS" in normalized
    assert "Completed at" in normalized
    assert "2026-02-05" in normalized
    assert "As of date" in normalized
    assert "2026-08-06" in normalized
    assert record["client"] not in normalized

    sidecar_path = out.with_name(out.name + ".provenance.json")
    assert sidecar_path.exists()
    with open(sidecar_path, encoding="utf-8") as f:
        loaded = json.load(f)

    assert loaded["completed_at"] == "2026-02-05"
    assert loaded["as_of_date"] == "2026-08-06"
    assert loaded["freshness_status"] == "STALE"
    assert loaded["freshness_reason"] == "OLDER_THAN_SIX_MONTHS"
    assert isinstance(loaded["content_hash"], str)
    assert len(loaded["content_hash"]) == 64
    assert re.fullmatch(r"[0-9a-f]{64}", loaded["content_hash"]) is not None
    assert record["client"] not in json.dumps(loaded)
